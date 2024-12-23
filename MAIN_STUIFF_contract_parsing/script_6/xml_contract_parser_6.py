import os
import re
import nltk
from docx import Document

# Import stop words and tokenizer
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# Download NLTK resources (run once)
nltk.download('stopwords')
nltk.download('punkt')

# Define the directory containing your contracts
contracts_dir = '/Users/amiyasekhar/CLM/contracts'  # Replace with your actual directory path

# Load English stop words
stop_words = set(stopwords.words('english'))

# Define clause types and associated keywords (expanded list)
clause_types = {
    'Acceptance': ['acceptance', 'accept', 'rejection', 'inspection'],
    'Amendments': ['amendments', 'modifications', 'changes', 'written consent', 'alteration'],
    'Assignment': ['assignment', 'assign', 'transfer', 'delegate', 'subcontract'],
    'Audit Rights': ['audit', 'inspection', 'records', 'books', 'examination', 'audit rights'],
    'Background Checks': ['background check', 'screening', 'criminal history', 'due diligence'],
    'Change of Control': ['change of control', 'merger', 'acquisition', 'sale of assets', 'ownership change'],
    'Compliance with Laws': ['compliance', 'laws', 'regulations', 'legal requirements', 'comply'],
    'Confidentiality': ['confidential', 'non-disclosure', 'secret', 'proprietary information', 'NDA'],
    'Conflict of Interest': ['conflict of interest', 'no conflict', 'duty of loyalty'],
    'Counterparts': ['counterparts', 'duplicate originals', 'facsimile', 'electronic signatures'],
    'Data Protection': ['data protection', 'personal data', 'GDPR', 'privacy', 'data security'],
    'Delivery': ['delivery', 'shipment', 'transportation', 'FOB', 'shipping terms'],
    'Dispute Resolution': ['dispute resolution', 'arbitration', 'mediation', 'litigation', 'court', 'disputes'],
    'Entire Agreement': ['entire agreement', 'whole agreement', 'complete agreement', 'supersedes', 'prior agreements'],
    'Environmental Compliance': ['environmental', 'pollution', 'hazardous materials', 'EPA', 'sustainability'],
    'Export Controls': ['export controls', 'export laws', 'trade compliance', 'embargo', 'sanctions'],
    'Force Majeure': ['force majeure', 'act of god', 'unforeseeable circumstances', 'beyond control', 'natural disaster'],
    'Further Assurances': ['further assurances', 'execute documents', 'take actions', 'necessary actions'],
    'Governing Currency': ['currency', 'payments in', 'USD', 'EUR', 'foreign exchange'],
    'Governing Law': ['governing law', 'jurisdiction', 'venue', 'applicable law', 'laws of'],
    'Governing Language': ['governing language', 'language of contract', 'translation'],
    'Headings': ['headings', 'titles', 'captions', 'section headings'],
    'Indemnification': ['indemnify', 'indemnification', 'hold harmless', 'defend', 'liabilities', 'claims'],
    'Insurance': ['insurance', 'insured', 'coverage', 'policy', 'liability insurance', 'certificate of insurance'],
    'Intellectual Property Rights': ['intellectual property', 'IP rights', 'copyright', 'trademark', 'patent', 'licensing'],
    'Limitation of Liability': ['limitation of liability', 'liability cap', 'maximum liability', 'exclude liability'],
    'Non-Compete': ['non-compete', 'competition', 'competitive', 'compete', 'restrictive covenant'],
    'Non-Discrimination': ['non-discrimination', 'equal opportunity', 'harassment', 'diversity', 'inclusion'],
    'Non-Solicitation': ['non-solicitation', 'solicit', 'poach', 'hire away', 'customers', 'employees'],
    'Notices': ['notices', 'notification', 'written notice', 'delivery', 'address', 'communication'],
    'Order of Precedence': ['order of precedence', 'conflicting terms', 'priority', 'hierarchy'],
    'Ownership': ['ownership', 'property rights', 'title', 'retain ownership', 'intellectual property'],
    'Payment Terms': ['payment terms', 'compensation', 'fees', 'charges', 'invoices', 'payment schedule'],
    'Privacy': ['privacy', 'confidentiality', 'personal information', 'data'],
    'Publicity': ['publicity', 'press release', 'announcement', 'marketing', 'advertising'],
    'Records Retention': ['records retention', 'document retention', 'recordkeeping', 'archive', 'storage'],
    'Relationship of Parties': ['relationship', 'independent contractor', 'agency', 'partnership', 'joint venture'],
    'Remedies': ['remedies', 'rights', 'cumulative', 'injunction', 'specific performance', 'damages'],
    'Representations and Warranties': ['representations', 'warranties', 'represent', 'warrant', 'no infringement', 'authority'],
    'Severability': ['severability', 'invalid provision', 'unenforceable', 'void provision', 'remain in effect'],
    'Subcontracting': ['subcontract', 'delegate', 'third parties', 'outsourcing', 'subcontractor'],
    'Survival': ['survival', 'continue in effect', 'remain in force', 'after termination', 'post-termination'],
    'Taxes': ['taxes', 'withholding', 'tax liability', 'tax obligations', 'taxes payable'],
    'Termination': ['terminate', 'termination', 'end of agreement', 'expiration', 'notice of termination'],
    'Termination for Cause': ['termination for cause', 'breach', 'default', 'failure to perform', 'material breach'],
    'Termination for Convenience': ['termination for convenience', 'without cause', 'no fault', 'at any time'],
    'Third Party Beneficiaries': ['third party beneficiaries', 'no third party rights', 'beneficiaries'],
    'Time is of the Essence': ['time is of the essence', 'timely', 'deadline', 'punctuality'],
    'Waiver': ['waiver', 'waive', 'failure to enforce', 'no waiver', 'rights waiver'],
    'Warranty': ['warranty', 'guarantee', 'defects', 'warrant', 'quality assurance'],
    'Compliance with Anti-Bribery Laws': ['anti-bribery', 'anti-corruption', 'FCPA', 'UK Bribery Act', 'bribes'],
    'Ethics and Compliance': ['ethics', 'code of conduct', 'ethical standards', 'compliance program'],
    'Set-Off': ['set-off', 'offset', 'deduct', 'withhold payment', 'counterclaim'],
    'Security': ['security', 'secure', 'safeguards', 'protection', 'unauthorized access', 'breach'],
    'Confidentiality and Non-Disclosure': ['confidentiality', 'non-disclosure', 'NDA', 'confidential information'],
    'Governing Regulations': ['governing regulations', 'compliance', 'statutes', 'legal compliance'],
    'Liens': ['lien', 'encumbrance', 'security interest', 'claim against property'],
    'No Partnership': ['no partnership', 'no joint venture', 'independent parties', 'agency'],
    'Non-Exclusive': ['non-exclusive', 'exclusive rights', 'sole rights', 'exclusivity'],
    'Patent Rights': ['patent', 'patented', 'patent rights', 'invention'],
    'Performance Standards': ['performance standards', 'service levels', 'SLAs', 'quality standards'],
    'Retention of Title': ['retention of title', 'title retention', 'ownership until payment'],
    'Subrogation': ['subrogation', 'rights of subrogation', 'insurance recovery'],
    'Technical Assistance': ['technical assistance', 'support', 'maintenance', 'service'],
    'Trade Secrets': ['trade secrets', 'confidential information', 'proprietary knowledge'],
    'User Obligations': ['user obligations', 'customer responsibilities', 'client duties'],
    'Warranty Disclaimer': ['warranty disclaimer', 'as is', 'no warranty', 'without warranty'],
    'Intellectual Property Indemnity': ['IP indemnity', 'intellectual property indemnification', 'IP claims'],
    'Limitation of Warranty': ['limitation of warranty', 'limited warranty', 'warranty period'],
    'Governing Standards': ['governing standards', 'industry standards', 'best practices'],
    'Product Recall': ['product recall', 'recall procedures', 'defective products'],
    'Inspection and Testing': ['inspection', 'testing', 'quality control', 'acceptance testing'],
    'Service Level Agreement': ['service level agreement', 'SLA', 'uptime', 'availability'],
    'Liquidated Damages': ['liquidated damages', 'pre-agreed damages', 'penalty', 'compensation'],
    'Governing Body': ['governing body', 'authority', 'regulatory agency'],
    'Health and Safety': ['health and safety', 'OSHA', 'workplace safety', 'safe practices'],
    'Intellectual Property Ownership': ['IP ownership', 'ownership rights', 'created IP', 'developed IP'],
    'Joint and Several Liability': ['joint liability', 'several liability', 'collective responsibility'],
    'Notice Period': ['notice period', 'advance notice', 'notification time'],
    'Obligations of the Parties': ['obligations', 'duties', 'responsibilities', 'commitments'],
    'Pricing and Fees': ['pricing', 'fees', 'charges', 'costs', 'rates'],
    'Quality Assurance': ['quality assurance', 'QA', 'quality control', 'standards'],
    'Regulatory Compliance': ['regulatory compliance', 'regulations', 'legal requirements'],
    'Rescission': ['rescission', 'cancelation', 'voiding', 'nullification'],
    'Retention of Records': ['retention of records', 'document retention', 'recordkeeping'],
    'Service Continuity': ['service continuity', 'business continuity', 'disaster recovery'],
    'Supplier Diversity': ['supplier diversity', 'minority-owned', 'diverse suppliers'],
    'Termination Assistance': ['termination assistance', 'transition services', 'exit assistance'],
    'Use of Subcontractors': ['subcontractors', 'third parties', 'outsourcing'],
    'Work Product': ['work product', 'deliverables', 'outputs', 'results'],
}

def classify_paragraph(paragraph_text):
    # Tokenize the content
    word_tokens = word_tokenize(paragraph_text)
    # Remove stop words
    filtered_words = [word for word in word_tokens if word.lower() not in stop_words]
    # Reconstruct the filtered content
    filtered_content = ' '.join(filtered_words)
    # Initialize a list to collect matching clause types
    matched_clause_types = []
    # Check for keywords to classify the paragraph
    for clause_type, keywords in clause_types.items():
        for keyword in keywords:
            if keyword.lower() in filtered_content.lower():
                matched_clause_types.append(clause_type)
                break  # Avoid duplicate entries if multiple keywords match
    # Remove duplicates (if any) and return the list
    if matched_clause_types:
        return list(set(matched_clause_types))
    else:
        return ['Other']

# Iterate over all files in the directory
for filename in os.listdir(contracts_dir):
    if filename.endswith('.docx'):
        file_path = os.path.join(contracts_dir, filename)
        print(f"Processing file: {filename}")

        # Load the Word document
        doc = Document(file_path)

        # Initialize variables
        paragraphs = []
        current_paragraph = ''

        # Process each paragraph in the document
        with open(f"/Users/amiyasekhar/CLM/MAIN_STUIFF_contract_parsing/results_6/{filename}.txt", 'w') as cf:
            num_paragraphs = len(doc.paragraphs)
            idx_para = 0
            while idx_para < num_paragraphs:
                para = doc.paragraphs[idx_para]
                text = para.text.strip()
                if text:
                    # If current_paragraph is empty, start a new paragraph
                    if not current_paragraph:
                        current_paragraph = text
                    else:
                        # Check if this paragraph starts with a lowercase letter
                        if text[0].islower():
                            # Continue the current paragraph
                            current_paragraph += ' ' + text
                        else:
                            # Treat as a new paragraph
                            paragraphs.append(current_paragraph)
                            current_paragraph = text
                else:
                    # Empty paragraph
                    # Peek at the next non-empty paragraph to decide
                    peek_idx = idx_para + 1
                    while peek_idx < num_paragraphs and not doc.paragraphs[peek_idx].text.strip():
                        peek_idx += 1
                    if peek_idx < num_paragraphs:
                        next_para_text = doc.paragraphs[peek_idx].text.strip()
                        if next_para_text and next_para_text[0].islower():
                            # Next paragraph starts with lowercase, continue current paragraph
                            pass
                        else:
                            # Next paragraph starts with uppercase or is empty, end current paragraph
                            if current_paragraph:
                                paragraphs.append(current_paragraph)
                                current_paragraph = ''
                    else:
                        # End of document
                        if current_paragraph:
                            paragraphs.append(current_paragraph)
                            current_paragraph = ''
                idx_para += 1

            # Add the last paragraph if it exists
            if current_paragraph:
                paragraphs.append(current_paragraph)

            # Classify each paragraph
            for idx, paragraph_text in enumerate(paragraphs, start=1):
                paragraph_types = classify_paragraph(paragraph_text)
                types_str = ', '.join(paragraph_types)
                cf.write(f"Paragraph {idx}:\n")
                cf.write(f"Type(s): {types_str}\n")
                cf.write(f"Content: {paragraph_text}\n")
                cf.write('-' * 80 + "\n")

            cf.write('\n' + '=' * 100 + '\n')
